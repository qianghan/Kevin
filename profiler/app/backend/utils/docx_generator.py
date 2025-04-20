"""
DOCX generation utilities.

This module provides functions for generating DOCX documents.
"""

import io
from typing import Dict, Any, Optional, Union
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(
    content: Union[str, bytes, Dict[str, Any]],
    metadata: Optional[Dict[str, Any]] = None,
    title: Optional[str] = None,
    author: Optional[str] = None,
    subject: Optional[str] = None,
    keywords: Optional[str] = None
) -> bytes:
    """
    Generate a DOCX document from content.
    
    Args:
        content: Content to convert to DOCX (text, HTML, or document data)
        metadata: Optional metadata to include in the DOCX
        title: Optional document title
        author: Optional document author
        subject: Optional document subject
        keywords: Optional document keywords
        
    Returns:
        DOCX document as bytes
    """
    # Create a new Document
    doc = Document()
    
    # Set document properties
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author
    if subject:
        doc.core_properties.subject = subject
    if keywords:
        doc.core_properties.keywords = keywords
    
    # Add title if provided
    if title:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing
    
    # Handle different content types
    if isinstance(content, str):
        # Check if content is HTML
        if content.strip().startswith('<'):
            # For now, just add as plain text
            # TODO: Add HTML parsing and conversion
            doc.add_paragraph(content)
        else:
            # Plain text
            doc.add_paragraph(content)
    elif isinstance(content, bytes):
        # Try to decode as text
        try:
            text_content = content.decode('utf-8')
            doc.add_paragraph(text_content)
        except UnicodeDecodeError:
            # If not text, add as binary data note
            doc.add_paragraph("Binary content cannot be displayed in DOCX format")
    elif isinstance(content, dict):
        # Handle structured document data
        if 'title' in content:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(content['title'])
            title_run.bold = True
            title_run.font.size = Pt(14)
            doc.add_paragraph()
        
        if 'content' in content:
            if isinstance(content['content'], str):
                doc.add_paragraph(content['content'])
            elif isinstance(content['content'], list):
                for item in content['content']:
                    doc.add_paragraph(str(item))
    
    # Add metadata if provided
    if metadata:
        doc.add_paragraph("\nMetadata:")
        for key, value in metadata.items():
            doc.add_paragraph(f"{key}: {value}")
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue() 