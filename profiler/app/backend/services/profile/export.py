from typing import Tuple, BinaryIO, Optional, Dict, Any
import os
import json
from datetime import datetime
from io import BytesIO
from pathlib import Path
import logging
from jinja2 import Environment, FileSystemLoader
import pdfkit
import weasyprint
from docx import Document
from docx.shared import Inches
import markdown
import yaml
import csv
import zipfile

from ..interfaces import ProfileExportInterface, ProfileExportError
from ..models import Profile, ProfileTemplate
from ..storage import ProfileStorage
from ..utils import format_date, sanitize_filename

logger = logging.getLogger(__name__)

class ProfileExportService(ProfileExportInterface):
    """Implementation of the profile export service."""
    
    def __init__(self, storage: ProfileStorage):
        self.storage = storage
        self.templates_dir = Path(__file__).parent / "templates"
        self.temp_dir = Path("/tmp/profile_exports")
        self.temp_dir.mkdir(exist_ok=True)
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
        
    async def initialize(self) -> None:
        """Initialize the export service."""
        logger.info("Initializing profile export service")
        # Ensure template directory exists
        self.templates_dir.mkdir(exist_ok=True)
        # Create default templates if they don't exist
        await self._create_default_templates()
        
    async def shutdown(self) -> None:
        """Shutdown the export service."""
        logger.info("Shutting down profile export service")
        # Clean up temporary files
        for file in self.temp_dir.glob("*"):
            try:
                file.unlink()
            except Exception as e:
                logger.warning(f"Failed to delete temporary file {file}: {e}")
                
    async def _create_default_templates(self) -> None:
        """Create default export templates if they don't exist."""
        default_templates = {
            "resume": {
                "name": "Default Resume Template",
                "description": "Professional resume template",
                "format": "pdf",
                "template": "resume_template.html"
            },
            "portfolio": {
                "name": "Default Portfolio Template",
                "description": "Professional portfolio template",
                "format": "html",
                "template": "portfolio_template.html"
            }
        }
        
        for template_id, template_data in default_templates.items():
            template_path = self.templates_dir / f"{template_id}.json"
            if not template_path.exists():
                with open(template_path, "w") as f:
                    json.dump(template_data, f, indent=2)
                    
    async def _get_template(self, template_id: str) -> Dict[str, Any]:
        """Get template configuration by ID."""
        template_path = self.templates_dir / f"{template_id}.json"
        if not template_path.exists():
            raise ProfileExportError(f"Template {template_id} not found")
            
        with open(template_path) as f:
            return json.load(f)
            
    async def _prepare_profile_data(self, profile: Profile) -> Dict[str, Any]:
        """Prepare profile data for export."""
        return {
            "personal_info": profile.personal_info,
            "education": profile.education,
            "experience": profile.experience,
            "skills": profile.skills,
            "projects": profile.projects,
            "certifications": profile.certifications,
            "languages": profile.languages,
            "interests": profile.interests,
            "references": profile.references,
            "metadata": {
                "export_date": format_date(datetime.now()),
                "profile_version": profile.version
            }
        } 

    async def export_profile(
        self,
        profile: Profile,
        format: str,
        template_id: Optional[str] = None,
        options: Optional[Dict[str, Any]] = None
    ) -> Tuple[str, BinaryIO]:
        """Export profile in specified format."""
        try:
            # Prepare profile data
            profile_data = await self._prepare_profile_data(profile)
            
            # Get template if specified
            template = None
            if template_id:
                template = await self._get_template(template_id)
                
            # Generate export based on format
            if format == "pdf":
                return await self._export_pdf(profile_data, template, options)
            elif format == "docx":
                return await self._export_docx(profile_data, template, options)
            elif format == "html":
                return await self._export_html(profile_data, template, options)
            elif format == "markdown":
                return await self._export_markdown(profile_data, template, options)
            elif format == "json":
                return await self._export_json(profile_data, template, options)
            elif format == "yaml":
                return await self._export_yaml(profile_data, template, options)
            elif format == "csv":
                return await self._export_csv(profile_data, template, options)
            else:
                raise ProfileExportError(f"Unsupported export format: {format}")
                
        except Exception as e:
            logger.error(f"Error exporting profile: {e}")
            raise ProfileExportError(f"Failed to export profile: {str(e)}")
            
    async def _export_pdf(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as PDF."""
        try:
            # Render HTML template
            html_template = self.jinja_env.get_template("pdf_template.html")
            html_content = html_template.render(profile=profile_data)
            
            # Convert to PDF
            pdf_options = {
                'page-size': 'A4',
                'margin-top': '20mm',
                'margin-right': '20mm',
                'margin-bottom': '20mm',
                'margin-left': '20mm',
                'encoding': 'UTF-8',
                'no-outline': None
            }
            
            if options:
                pdf_options.update(options)
                
            pdf_content = pdfkit.from_string(html_content, False, options=pdf_options)
            
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.pdf"
            filename = sanitize_filename(filename)
            
            return filename, BytesIO(pdf_content)
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise ProfileExportError(f"Failed to generate PDF: {str(e)}")
            
    async def _export_docx(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as DOCX."""
        try:
            doc = Document()
            
            # Add title
            doc.add_heading(f"{profile_data['personal_info']['name']} - Professional Profile", 0)
            
            # Add sections
            for section, data in profile_data.items():
                if section == "metadata":
                    continue
                    
                doc.add_heading(section.replace("_", " ").title(), level=1)
                
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict):
                            for key, value in item.items():
                                doc.add_paragraph(f"{key}: {value}")
                        else:
                            doc.add_paragraph(str(item))
                else:
                    doc.add_paragraph(str(data))
                    
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.docx"
            filename = sanitize_filename(filename)
            
            return filename, buffer
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise ProfileExportError(f"Failed to generate DOCX: {str(e)}")
            
    async def _export_html(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as HTML."""
        try:
            # Render HTML template
            html_template = self.jinja_env.get_template("html_template.html")
            html_content = html_template.render(profile=profile_data)
            
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.html"
            filename = sanitize_filename(filename)
            
            return filename, BytesIO(html_content.encode())
            
        except Exception as e:
            logger.error(f"Error generating HTML: {e}")
            raise ProfileExportError(f"Failed to generate HTML: {str(e)}")
            
    async def _export_markdown(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as Markdown."""
        try:
            markdown_content = []
            
            # Add title
            markdown_content.append(f"# {profile_data['personal_info']['name']} - Professional Profile\n")
            
            # Add sections
            for section, data in profile_data.items():
                if section == "metadata":
                    continue
                    
                markdown_content.append(f"## {section.replace('_', ' ').title()}\n")
                
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict):
                            for key, value in item.items():
                                markdown_content.append(f"- **{key}**: {value}")
                        else:
                            markdown_content.append(f"- {item}")
                else:
                    markdown_content.append(str(data))
                    
                markdown_content.append("")
                
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.md"
            filename = sanitize_filename(filename)
            
            return filename, BytesIO("\n".join(markdown_content).encode())
            
        except Exception as e:
            logger.error(f"Error generating Markdown: {e}")
            raise ProfileExportError(f"Failed to generate Markdown: {str(e)}")
            
    async def _export_json(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as JSON."""
        try:
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.json"
            filename = sanitize_filename(filename)
            
            return filename, BytesIO(json.dumps(profile_data, indent=2).encode())
            
        except Exception as e:
            logger.error(f"Error generating JSON: {e}")
            raise ProfileExportError(f"Failed to generate JSON: {str(e)}")
            
    async def _export_yaml(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as YAML."""
        try:
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.yaml"
            filename = sanitize_filename(filename)
            
            return filename, BytesIO(yaml.dump(profile_data, default_flow_style=False).encode())
            
        except Exception as e:
            logger.error(f"Error generating YAML: {e}")
            raise ProfileExportError(f"Failed to generate YAML: {str(e)}")
            
    async def _export_csv(
        self,
        profile_data: Dict[str, Any],
        template: Optional[Dict[str, Any]],
        options: Optional[Dict[str, Any]]
    ) -> Tuple[str, BinaryIO]:
        """Export profile as CSV."""
        try:
            buffer = BytesIO()
            writer = csv.writer(buffer)
            
            # Write header
            writer.writerow(["Section", "Field", "Value"])
            
            # Write data
            for section, data in profile_data.items():
                if section == "metadata":
                    continue
                    
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict):
                            for key, value in item.items():
                                writer.writerow([section, key, value])
                        else:
                            writer.writerow([section, "", item])
                else:
                    writer.writerow([section, "", data])
                    
            buffer.seek(0)
            
            # Create filename
            filename = f"profile_{profile_data['personal_info']['name']}_{datetime.now().strftime('%Y%m%d')}.csv"
            filename = sanitize_filename(filename)
            
            return filename, buffer
            
        except Exception as e:
            logger.error(f"Error generating CSV: {e}")
            raise ProfileExportError(f"Failed to generate CSV: {str(e)}")
            
    async def export_profile_archive(
        self,
        profile: Profile,
        formats: list[str],
        options: Optional[Dict[str, Any]] = None
    ) -> Tuple[str, BinaryIO]:
        """Export profile in multiple formats as a ZIP archive."""
        try:
            # Create ZIP file
            buffer = BytesIO()
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Export in each format
                for format in formats:
                    filename, content = await self.export_profile(profile, format, options=options)
                    zip_file.writestr(filename, content.getvalue())
                    
            buffer.seek(0)
            
            # Create archive filename
            archive_filename = f"profile_{profile.personal_info['name']}_{datetime.now().strftime('%Y%m%d')}.zip"
            archive_filename = sanitize_filename(archive_filename)
            
            return archive_filename, buffer
            
        except Exception as e:
            logger.error(f"Error generating profile archive: {e}")
            raise ProfileExportError(f"Failed to generate profile archive: {str(e)}") 