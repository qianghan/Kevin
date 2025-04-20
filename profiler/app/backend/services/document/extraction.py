"""
Document metadata extraction service.

This module provides functionality for extracting metadata from different document types.
"""

import os
import re
import json
import tempfile
from typing import Dict, List, Any, Optional, BinaryIO, Tuple
from datetime import datetime

from app.backend.utils.logging import get_logger
from app.backend.utils.errors import StorageError, ValidationError
from .models import Document

logger = get_logger(__name__)


class DocumentMetadataExtractor:
    """Base class for document metadata extractors."""
    
    @classmethod
    def get_extractor(cls, mime_type: str) -> 'DocumentMetadataExtractor':
        """
        Get an appropriate extractor for the given MIME type.
        
        Args:
            mime_type: MIME type of the document
            
        Returns:
            DocumentMetadataExtractor instance
        """
        # Map MIME types to extractor classes
        extractors = {
            # PDF files
            "application/pdf": PDFMetadataExtractor,
            
            # Office documents
            "application/msword": OfficeMetadataExtractor,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": OfficeMetadataExtractor,
            "application/vnd.ms-excel": OfficeMetadataExtractor,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": OfficeMetadataExtractor,
            "application/vnd.ms-powerpoint": OfficeMetadataExtractor,
            "application/vnd.openxmlformats-officedocument.presentationml.presentation": OfficeMetadataExtractor,
            
            # Text files
            "text/plain": TextMetadataExtractor,
            "text/html": HTMLMetadataExtractor,
            "text/xml": XMLMetadataExtractor,
            "application/json": JSONMetadataExtractor,
            "text/markdown": MarkdownMetadataExtractor,
            
            # Images
            "image/jpeg": ImageMetadataExtractor,
            "image/png": ImageMetadataExtractor,
            "image/gif": ImageMetadataExtractor,
            "image/tiff": ImageMetadataExtractor
        }
        
        # Get extractor for MIME type, or use default
        extractor_class = extractors.get(mime_type, DefaultMetadataExtractor)
        return extractor_class()
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from a document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
            
        Raises:
            StorageError: If metadata cannot be extracted
        """
        # Default implementation that should be overridden
        return {
            "extraction_time": datetime.utcnow().isoformat(),
            "extractor": self.__class__.__name__
        }
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        _, ext = os.path.splitext(filename)
        return ext.lower()


class DefaultMetadataExtractor(DocumentMetadataExtractor):
    """Default metadata extractor for unsupported document types."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract basic metadata from a document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            # Get file size
            file_content.seek(0, os.SEEK_END)
            file_size = file_content.tell()
            
            # Read first few bytes for a simple file signature
            file_content.seek(0)
            signature = file_content.read(16).hex()
            
            # Restore original position
            file_content.seek(pos)
            
            return {
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_size": file_size,
                "file_extension": self._get_file_extension(filename),
                "file_signature": signature
            }
            
        except Exception as e:
            logger.error(f"Failed to extract metadata from {filename}: {str(e)}")
            
            # Restore original position
            file_content.seek(pos)
            
            return {
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "extraction_error": str(e)
            }


class PDFMetadataExtractor(DocumentMetadataExtractor):
    """Metadata extractor for PDF documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from a PDF document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            try:
                # Try to import PyPDF2, fall back to default if not available
                import PyPDF2
            except ImportError:
                logger.warning("PyPDF2 not installed; using default metadata extraction.")
                return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
            # Create PDF reader
            pdf_reader = PyPDF2.PdfReader(file_content)
            
            # Get document info
            info = pdf_reader.metadata
            if info:
                # Convert PDF info to dictionary
                metadata = {
                    "title": info.get('/Title', ''),
                    "author": info.get('/Author', ''),
                    "subject": info.get('/Subject', ''),
                    "keywords": info.get('/Keywords', ''),
                    "creator": info.get('/Creator', ''),
                    "producer": info.get('/Producer', ''),
                    "creation_date": info.get('/CreationDate', ''),
                    "modification_date": info.get('/ModDate', '')
                }
            else:
                metadata = {}
            
            # Add page count
            metadata["page_count"] = len(pdf_reader.pages)
            
            # Get text from first page for content preview
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                try:
                    text = first_page.extract_text()
                    if text:
                        # Limit preview to ~1000 chars
                        metadata["content_preview"] = text[:1000]
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF: {str(e)}")
            
            # Add standard metadata
            metadata.update({
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_extension": self._get_file_extension(filename)
            })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract PDF metadata from {filename}: {str(e)}")
            
            # Fall back to default extractor
            return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class OfficeMetadataExtractor(DocumentMetadataExtractor):
    """Metadata extractor for Microsoft Office documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from Office documents.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            try:
                # Try to import python-docx for DOCX files
                import docx
                # Try to import openpyxl for XLSX files
                import openpyxl
            except ImportError:
                logger.warning("Office document libraries not installed; using default metadata extraction.")
                return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
            # Create temporary file to work with
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                file_content.seek(0)
                temp_file.write(file_content.read())
                temp_path = temp_file.name
            
            try:
                ext = self._get_file_extension(filename)
                
                if ext == '.docx':
                    # Extract DOCX metadata
                    doc = docx.Document(temp_path)
                    core_props = doc.core_properties
                    
                    metadata = {
                        "title": core_props.title or '',
                        "author": core_props.author or '',
                        "comments": core_props.comments or '',
                        "keywords": core_props.keywords or '',
                        "last_modified_by": core_props.last_modified_by or '',
                        "created": core_props.created.isoformat() if core_props.created else '',
                        "modified": core_props.modified.isoformat() if core_props.modified else '',
                        "paragraph_count": len(doc.paragraphs),
                        "section_count": len(doc.sections),
                        "content_preview": doc.paragraphs[0].text[:1000] if doc.paragraphs else ''
                    }
                    
                elif ext == '.xlsx':
                    # Extract XLSX metadata
                    wb = openpyxl.load_workbook(temp_path, read_only=True)
                    props = wb.properties
                    
                    metadata = {
                        "title": props.title or '',
                        "creator": props.creator or '',
                        "subject": props.subject or '',
                        "description": props.description or '',
                        "keywords": props.keywords or '',
                        "category": props.category or '',
                        "created": props.created.isoformat() if props.created else '',
                        "modified": props.modified.isoformat() if props.modified else '',
                        "sheet_names": wb.sheetnames,
                        "sheet_count": len(wb.sheetnames)
                    }
                    
                else:
                    # Unsupported Office format
                    return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
                
                # Add standard metadata
                metadata.update({
                    "extraction_time": datetime.utcnow().isoformat(),
                    "extractor": self.__class__.__name__,
                    "file_extension": ext
                })
                
                return metadata
                
            finally:
                # Delete temporary file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            
        except Exception as e:
            logger.error(f"Failed to extract Office metadata from {filename}: {str(e)}")
            
            # Fall back to default extractor
            return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class TextMetadataExtractor(DocumentMetadataExtractor):
    """Metadata extractor for plain text documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from a text document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            # Read content
            file_content.seek(0)
            content = file_content.read()
            
            try:
                # Try to decode as UTF-8
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Fall back to latin-1 for binary files
                text = content.decode('latin-1')
            
            # Count lines, words, and characters
            lines = text.splitlines()
            line_count = len(lines)
            word_count = len(text.split())
            char_count = len(text)
            
            # Create content preview
            content_preview = text[:1000] if text else ''
            
            # Get common headers from text files (e.g., title, author)
            headers = {}
            for line in lines[:20]:  # Check first 20 lines
                # Look for common headers
                header_match = re.match(r'^([A-Za-z\-]+):\s*(.+)$', line)
                if header_match:
                    key, value = header_match.groups()
                    headers[key.lower()] = value
            
            metadata = {
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_extension": self._get_file_extension(filename),
                "line_count": line_count,
                "word_count": word_count,
                "character_count": char_count,
                "content_preview": content_preview,
                "headers": headers
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract text metadata from {filename}: {str(e)}")
            
            # Fall back to default extractor
            return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class HTMLMetadataExtractor(TextMetadataExtractor):
    """Metadata extractor for HTML documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from an HTML document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            try:
                # Try to import BeautifulSoup, fall back to default if not available
                from bs4 import BeautifulSoup
            except ImportError:
                logger.warning("BeautifulSoup not installed; using text metadata extraction.")
                return await super().extract_metadata(file_content, filename)
            
            # Read content
            file_content.seek(0)
            content = file_content.read()
            
            try:
                # Try to decode as UTF-8
                html_text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Fall back to latin-1 for binary files
                html_text = content.decode('latin-1')
            
            # Parse HTML
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Get metadata from head
            metadata = {}
            
            # Get title
            title_tag = soup.find('title')
            if title_tag:
                metadata["title"] = title_tag.string.strip() if title_tag.string else ''
            
            # Get meta tags
            for meta_tag in soup.find_all('meta'):
                name = meta_tag.get('name', meta_tag.get('property', ''))
                content = meta_tag.get('content', '')
                if name and content:
                    metadata[name.lower()] = content
            
            # Get text content for preview
            body_text = soup.get_text(separator=' ', strip=True)
            metadata["content_preview"] = body_text[:1000] if body_text else ''
            
            # Count elements
            metadata["link_count"] = len(soup.find_all('a'))
            metadata["image_count"] = len(soup.find_all('img'))
            metadata["heading_count"] = len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']))
            
            # Add standard metadata
            metadata.update({
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_extension": self._get_file_extension(filename)
            })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract HTML metadata from {filename}: {str(e)}")
            
            # Fall back to text extractor
            return await super().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class XMLMetadataExtractor(TextMetadataExtractor):
    """Metadata extractor for XML documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from an XML document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            try:
                # Try to import ElementTree, fall back to default if not available
                import xml.etree.ElementTree as ET
            except ImportError:
                logger.warning("ElementTree not available; using text metadata extraction.")
                return await super().extract_metadata(file_content, filename)
            
            # Read content
            file_content.seek(0)
            
            # Parse XML
            tree = ET.parse(file_content)
            root = tree.getroot()
            
            # Get root element and namespaces
            root_tag = root.tag
            nsmap = root.nsmap if hasattr(root, 'nsmap') else {}
            
            # Count elements
            element_count = sum(1 for _ in root.iter())
            
            # Get attributes from root
            root_attrs = {k: v for k, v in root.attrib.items()}
            
            metadata = {
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_extension": self._get_file_extension(filename),
                "root_element": root_tag,
                "namespaces": nsmap,
                "element_count": element_count,
                "root_attributes": root_attrs
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract XML metadata from {filename}: {str(e)}")
            
            # Fall back to text extractor
            return await super().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class JSONMetadataExtractor(TextMetadataExtractor):
    """Metadata extractor for JSON documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from a JSON document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            # Read content
            file_content.seek(0)
            content = file_content.read()
            
            try:
                # Try to decode as UTF-8
                json_text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Fall back to latin-1 for binary files
                json_text = content.decode('latin-1')
            
            # Parse JSON
            data = json.loads(json_text)
            
            # Analyze structure
            if isinstance(data, dict):
                structure_type = "object"
                top_level_keys = list(data.keys())
                key_count = len(top_level_keys)
            elif isinstance(data, list):
                structure_type = "array"
                key_count = 0
                top_level_keys = []
                array_length = len(data)
                
                # Get sample of first item if it's an object
                if array_length > 0 and isinstance(data[0], dict):
                    top_level_keys = list(data[0].keys())
                    key_count = len(top_level_keys)
            else:
                structure_type = "primitive"
                top_level_keys = []
                key_count = 0
            
            metadata = {
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_extension": self._get_file_extension(filename),
                "structure_type": structure_type,
                "top_level_keys": top_level_keys[:20],  # Limit to first 20 keys
                "key_count": key_count
            }
            
            # Add array length if applicable
            if structure_type == "array":
                metadata["array_length"] = array_length
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract JSON metadata from {filename}: {str(e)}")
            
            # Fall back to text extractor
            return await super().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class MarkdownMetadataExtractor(TextMetadataExtractor):
    """Metadata extractor for Markdown documents."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from a Markdown document.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            # Read content
            file_content.seek(0)
            content = file_content.read()
            
            try:
                # Try to decode as UTF-8
                md_text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Fall back to latin-1 for binary files
                md_text = content.decode('latin-1')
            
            # Extract front matter if present
            front_matter = {}
            if md_text.startswith('---'):
                # YAML front matter
                parts = md_text.split('---', 2)
                if len(parts) >= 3:
                    try:
                        import yaml
                        front_matter = yaml.safe_load(parts[1])
                    except (ImportError, yaml.YAMLError):
                        # Handle missing yaml package or parsing errors
                        front_matter_text = parts[1]
                        # Primitive front matter parsing
                        for line in front_matter_text.splitlines():
                            if ':' in line:
                                key, value = line.split(':', 1)
                                front_matter[key.strip()] = value.strip()
                    
                    # Remove front matter from content
                    md_text = parts[2]
            
            # Count various elements
            heading_count = len(re.findall(r'^#+\s+.+$', md_text, re.MULTILINE))
            link_count = len(re.findall(r'\[.*?\]\(.*?\)', md_text))
            code_block_count = len(re.findall(r'```.*?```', md_text, re.DOTALL))
            list_item_count = len(re.findall(r'^[\*\-\+]\s+.+$', md_text, re.MULTILINE))
            
            # Get title from first heading
            title_match = re.search(r'^#\s+(.+)$', md_text, re.MULTILINE)
            title = title_match.group(1) if title_match else ''
            
            # Get content preview
            lines = md_text.splitlines()
            content_lines = [l for l in lines if l.strip() and not l.startswith('#')]
            content_preview = ' '.join(content_lines[:5]) if content_lines else ''
            
            metadata = {
                "extraction_time": datetime.utcnow().isoformat(),
                "extractor": self.__class__.__name__,
                "file_extension": self._get_file_extension(filename),
                "front_matter": front_matter,
                "title": title,
                "heading_count": heading_count,
                "link_count": link_count,
                "code_block_count": code_block_count,
                "list_item_count": list_item_count,
                "content_preview": content_preview[:1000]
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract Markdown metadata from {filename}: {str(e)}")
            
            # Fall back to text extractor
            return await super().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class ImageMetadataExtractor(DocumentMetadataExtractor):
    """Metadata extractor for image files."""
    
    async def extract_metadata(self, file_content: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from an image file.
        
        Args:
            file_content: Document content
            filename: Original filename
            
        Returns:
            Extracted metadata
        """
        # Get current position to restore later
        pos = file_content.tell()
        
        try:
            try:
                # Try to import Pillow, fall back to default if not available
                from PIL import Image
                from PIL.ExifTags import TAGS
            except ImportError:
                logger.warning("Pillow not installed; using default metadata extraction.")
                return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
            # Create temporary file to work with
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                file_content.seek(0)
                temp_file.write(file_content.read())
                temp_path = temp_file.name
            
            try:
                # Open image
                img = Image.open(temp_path)
                
                # Get basic info
                metadata = {
                    "format": img.format,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height,
                    "size": f"{img.width} x {img.height} pixels"
                }
                
                # Get EXIF data if available
                exif_data = {}
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        for tag_id, value in exif.items():
                            tag = TAGS.get(tag_id, tag_id)
                            exif_data[tag] = str(value)
                
                # Add EXIF data if found
                if exif_data:
                    # Filter out binary data and limit to common tags
                    filtered_exif = {}
                    common_tags = [
                        'Make', 'Model', 'DateTime', 'DateTimeOriginal', 'DateTimeDigitized',
                        'XResolution', 'YResolution', 'ExposureTime', 'FNumber', 'ISOSpeedRatings',
                        'FocalLength', 'Flash', 'Orientation', 'Software'
                    ]
                    
                    for tag, value in exif_data.items():
                        if tag in common_tags:
                            filtered_exif[tag] = value
                    
                    metadata["exif"] = filtered_exif
                
                # Add standard metadata
                metadata.update({
                    "extraction_time": datetime.utcnow().isoformat(),
                    "extractor": self.__class__.__name__,
                    "file_extension": self._get_file_extension(filename)
                })
                
                return metadata
                
            finally:
                # Delete temporary file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            
        except Exception as e:
            logger.error(f"Failed to extract image metadata from {filename}: {str(e)}")
            
            # Fall back to default extractor
            return await DefaultMetadataExtractor().extract_metadata(file_content, filename)
            
        finally:
            # Restore original position
            file_content.seek(pos)


class DocumentMetadataExtractionService:
    """Service for extracting metadata from documents."""
    
    async def extract_metadata(self, document: Document, file_content: BinaryIO) -> Dict[str, Any]:
        """
        Extract metadata from a document.
        
        Args:
            document: Document to extract metadata from
            file_content: Document content
            
        Returns:
            Extracted metadata
            
        Raises:
            StorageError: If metadata cannot be extracted
        """
        try:
            # Get appropriate extractor based on MIME type
            extractor = DocumentMetadataExtractor.get_extractor(document.mime_type)
            
            # Extract metadata
            metadata = await extractor.extract_metadata(file_content, document.filename)
            
            # Add document info
            metadata.update({
                "document_id": document.document_id,
                "user_id": document.user_id,
                "filename": document.filename,
                "mime_type": document.mime_type,
                "file_size": document.file_size,
                "status": document.status
            })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract metadata from document {document.document_id}: {str(e)}")
            raise StorageError(f"Failed to extract document metadata: {str(e)}")
    
    async def update_document_with_metadata(self, 
                                          document: Document, 
                                          file_content: BinaryIO,
                                          document_repository) -> Document:
        """
        Extract metadata from a document and update the document with it.
        
        Args:
            document: Document to extract metadata from
            file_content: Document content
            document_repository: Repository to update document in
            
        Returns:
            Updated document
            
        Raises:
            StorageError: If metadata cannot be extracted or document cannot be updated
        """
        try:
            # Extract metadata
            metadata = await self.extract_metadata(document, file_content)
            
            # Update document metadata
            current_metadata = document.metadata or {}
            
            # Add metadata under 'extracted_metadata' key
            current_metadata["extracted_metadata"] = metadata
            
            # Update document
            document.metadata = current_metadata
            
            # Save updated document
            updated_document = await document_repository.save_document(document)
            
            return updated_document
            
        except Exception as e:
            logger.error(f"Failed to update document with metadata {document.document_id}: {str(e)}")
            raise StorageError(f"Failed to update document with metadata: {str(e)}") 